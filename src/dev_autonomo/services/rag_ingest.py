"""RAG ingest service — pipeline de extracao + chunking + PII filter + embed + upsert.

Bloco C do roadmap stack-knowledge. Suporta 3 modos de input:
- file_upload: PDF (pypdf), DOCX (python-docx), TXT/MD (direto)
- url_fetch: HTML extraido via trafilatura
- pasted_text: texto direto

Pipeline:
    raw_source -> extract_text() -> apply_pii_filter() -> chunk_text()
    -> generate_embeddings() (Voyage) -> upsert_to_qdrant() -> register_rag_source()

Idempotencia: source_hash (SHA-256 do texto extraido) + UNIQUE
constraint em rag_sources(collection_slug, source_hash) impede
duplicacao da mesma fonte.

Custos: Voyage code-3 ~US$ 0.18 por 1M tokens. Doc de 500 paginas
~50k tokens = ~US$ 0.01 por upload.
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
from dataclasses import dataclass
from typing import BinaryIO
from uuid import UUID, uuid4

from qdrant_client.models import PointStruct
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from dev_autonomo.common.enums import (
    RagSourceKind,
    RagSourceLicense,
    RagSourceQuality,
    RagSourceScope,
    RagSourceStatus,
)
from dev_autonomo.db.models import RagSource
from dev_autonomo.knowledge.qdrant_client import QdrantKnowledgeStore
from dev_autonomo.knowledge.voyage_client import VoyageEmbeddingClient

logger = logging.getLogger(__name__)


CHUNK_SIZE_CHARS = 3200  # ~800 tokens
CHUNK_OVERLAP_CHARS = 400  # ~100 tokens


# ---------------------------------------------------------------------------
# Tipos
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class IngestRequest:
    """Pedido de ingest de uma fonte."""

    collection_slug: str           # ex: stack_patterns:hybris
    kind: RagSourceKind
    scope: RagSourceScope
    license: RagSourceLicense
    source_quality: RagSourceQuality
    stack_version: str | None = None
    tags: list[str] | None = None
    source_uri: str | None = None  # path, URL, ou None pra pasted_text
    client_id: UUID | None = None  # obrigatorio quando scope=client_private
    uploaded_by_user_id: UUID | None = None


@dataclass(slots=True)
class IngestResult:
    """Resultado do pipeline."""

    rag_source_id: UUID
    status: RagSourceStatus
    indexed_chunks: int
    source_hash: str
    error_message: str | None = None


# ---------------------------------------------------------------------------
# Extracao
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extrai texto de PDF via pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("falha ao extrair pagina: %s", exc)
    return "\n\n".join(parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extrai texto de DOCX via python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Inclui texto de tabelas (simples — linha por linha).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Detecta tipo por extensao e extrai."""
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if name_lower.endswith((".docx", ".doc")):
        return extract_text_from_docx(file_bytes)
    if name_lower.endswith((".md", ".markdown", ".txt")):
        # Texto direto — decode UTF-8 com fallback.
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="replace")
    raise ValueError(
        f"extensao nao suportada: {filename}. Aceitos: .pdf .docx .md .txt"
    )


def extract_text_from_url(url: str) -> str:
    """Fetch + extracao via trafilatura."""
    import trafilatura

    downloaded = trafilatura.fetch_url(url)
    if downloaded is None:
        raise ValueError(f"falha ao baixar URL: {url}")
    text = trafilatura.extract(
        downloaded,
        include_comments=False,
        include_tables=True,
        deduplicate=True,
    )
    if not text:
        raise ValueError(f"trafilatura nao extraiu texto de: {url}")
    return text


# ---------------------------------------------------------------------------
# PII filter
# ---------------------------------------------------------------------------


# Regex agressivos pra remover/mascarar info sensivel antes de embed.
# - Emails completos
# - Bearer tokens / API keys com prefixos conhecidos
# - URLs com auth (https://user:pass@host)
# - IPs privados
_PII_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}"), "<email>"),
    (re.compile(r"(?:Bearer\s+)?[A-Za-z0-9_-]{40,}\b"), "<token>"),
    (re.compile(r"(?:sk|pk|gh[ps]|github_pat)_[A-Za-z0-9]{20,}"), "<api_key>"),
    (re.compile(r"https?://[^:\s/]+:[^@\s/]+@"), "https://<redacted>@"),
    (re.compile(r"\b(?:10|172\.(?:1[6-9]|2\d|3[01])|192\.168)\.\d{1,3}\.\d{1,3}\b"), "<private_ip>"),
]


def apply_pii_filter(text: str) -> str:
    """Aplica regex de PII em texto antes de embedar."""
    out = text
    for pattern, replacement in _PII_PATTERNS:
        out = pattern.sub(replacement, out)
    return out


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE_CHARS,
    overlap: int = CHUNK_OVERLAP_CHARS,
) -> list[str]:
    """Chunk simples por sliding window com overlap.

    Pra texto livre (PDF, MD, web). Pra codigo use chunker.py existente.
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


# ---------------------------------------------------------------------------
# Hash + collection helpers
# ---------------------------------------------------------------------------


def compute_source_hash(text: str) -> str:
    """SHA-256 do texto extraido (pos PII filter)."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def qdrant_collection_for_slug(collection_slug: str) -> str:
    """Mapeia collection_slug (rag_sources) pra nome real Qdrant.

    Convencao: dev_autonomo.{collection_slug com `:` virou `_`}.
    Qdrant nao aceita `:` em nomes de colecao.
    """
    safe = collection_slug.replace(":", "_")
    return f"dev_autonomo.{safe}"


# ---------------------------------------------------------------------------
# Ingest pipeline
# ---------------------------------------------------------------------------


async def ingest(
    session: AsyncSession,
    request: IngestRequest,
    raw_text: str | None = None,
    file_bytes: bytes | None = None,
    file_name: str | None = None,
    voyage_client: VoyageEmbeddingClient | None = None,
    qdrant_store: QdrantKnowledgeStore | None = None,
) -> IngestResult:
    """Orquestra pipeline completo. Caller deve commitar a session apos.

    Inputs validos por kind:
    - PASTED_TEXT: raw_text obrigatorio.
    - FILE_UPLOAD: file_bytes + file_name obrigatorios.
    - URL_FETCH: request.source_uri obrigatorio.
    - FEEDBACK_LOOP / DREAMING: raw_text obrigatorio.
    """
    rag_source = RagSource(
        id=uuid4(),
        collection_slug=request.collection_slug,
        kind=request.kind,
        source_uri=request.source_uri,
        source_hash="pending",  # preenchido apos extracao
        scope=request.scope,
        client_id=request.client_id,
        license=request.license,
        source_quality=request.source_quality,
        stack_version=request.stack_version,
        uploaded_by_user_id=request.uploaded_by_user_id,
        indexed_chunks=0,
        status=RagSourceStatus.PENDING,
        tags=request.tags or [],
    )
    session.add(rag_source)

    try:
        # 1. Extrair
        rag_source.status = RagSourceStatus.EXTRACTING
        await session.flush()

        if request.kind == RagSourceKind.PASTED_TEXT:
            if not raw_text:
                raise ValueError("PASTED_TEXT requer raw_text")
            text = raw_text
        elif request.kind == RagSourceKind.FILE_UPLOAD:
            if not file_bytes or not file_name:
                raise ValueError("FILE_UPLOAD requer file_bytes + file_name")
            text = extract_text_from_file(file_name, file_bytes)
        elif request.kind == RagSourceKind.URL_FETCH:
            if not request.source_uri:
                raise ValueError("URL_FETCH requer source_uri")
            text = extract_text_from_url(request.source_uri)
        elif request.kind in (RagSourceKind.FEEDBACK_LOOP, RagSourceKind.DREAMING):
            if not raw_text:
                raise ValueError(f"{request.kind.value} requer raw_text")
            text = raw_text
        else:
            raise ValueError(f"kind nao suportado: {request.kind}")

        if not text or len(text.strip()) < 50:
            raise ValueError(f"texto extraido vazio ou curto demais (<50 chars)")

        # 2. PII filter
        clean_text = apply_pii_filter(text)

        # 3. Hash + dedup check
        source_hash = compute_source_hash(clean_text)
        rag_source.source_hash = source_hash

        existing = (await session.execute(
            select(RagSource).where(
                RagSource.collection_slug == request.collection_slug,
                RagSource.source_hash == source_hash,
                RagSource.id != rag_source.id,
            )
        )).scalar_one_or_none()
        if existing is not None:
            # Fonte ja existe — invalida a nova entry e retorna a existente.
            await session.delete(rag_source)
            await session.flush()
            return IngestResult(
                rag_source_id=existing.id,
                status=existing.status,
                indexed_chunks=existing.indexed_chunks,
                source_hash=source_hash,
                error_message="ja existe (dedup por hash)",
            )

        # 4. Chunk
        chunks = chunk_text(clean_text)
        if not chunks:
            raise ValueError("chunker nao gerou chunks")

        # 5. Embed
        rag_source.status = RagSourceStatus.EMBEDDING
        await session.flush()

        voyage = voyage_client or VoyageEmbeddingClient()
        embed_result = await voyage.embed_documents(chunks)

        # 6. Upsert Qdrant
        qdrant = qdrant_store or QdrantKnowledgeStore()
        collection_name = qdrant_collection_for_slug(request.collection_slug)

        # ensure_collection precisa de partition+squad_id no helper atual; pra
        # stack_patterns usamos client direto.
        from qdrant_client.models import Distance, VectorParams

        await qdrant._client.recreate_collection if False else None  # placeholder noop
        # Cria se nao existe (idempotente).
        try:
            await qdrant._client.get_collection(collection_name)
        except Exception:
            await qdrant._client.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(size=len(embed_result.vectors[0]), distance=Distance.COSINE),
            )

        points = [
            PointStruct(
                id=str(uuid4()),
                vector=vector,
                payload={
                    "rag_source_id": str(rag_source.id),
                    "collection_slug": request.collection_slug,
                    "scope": request.scope.value,
                    "client_id": str(request.client_id) if request.client_id else None,
                    "license": request.license.value,
                    "source_quality": request.source_quality.value,
                    "stack_version": request.stack_version,
                    "source_uri": request.source_uri,
                    "tags": request.tags or [],
                    "chunk_index": idx,
                    "chunk_total": len(chunks),
                    "content": chunk[:4000],
                },
            )
            for idx, (chunk, vector) in enumerate(zip(chunks, embed_result.vectors, strict=True))
        ]
        await qdrant._client.upsert(collection_name=collection_name, points=points)

        rag_source.indexed_chunks = len(chunks)
        rag_source.status = RagSourceStatus.INDEXED
        await session.flush()

        return IngestResult(
            rag_source_id=rag_source.id,
            status=RagSourceStatus.INDEXED,
            indexed_chunks=len(chunks),
            source_hash=source_hash,
        )

    except Exception as exc:
        logger.exception("ingest falhou collection=%s: %s",
                         request.collection_slug, exc)
        rag_source.status = RagSourceStatus.FAILED
        rag_source.error_message = f"{type(exc).__name__}: {exc}"
        await session.flush()
        return IngestResult(
            rag_source_id=rag_source.id,
            status=RagSourceStatus.FAILED,
            indexed_chunks=0,
            source_hash=rag_source.source_hash,
            error_message=rag_source.error_message,
        )


async def delete_source(
    session: AsyncSession,
    source_id: UUID,
    qdrant_store: QdrantKnowledgeStore | None = None,
) -> bool:
    """Remove rag_source e seus chunks no Qdrant. Retorna True se removeu."""
    source = await session.get(RagSource, source_id)
    if source is None:
        return False

    qdrant = qdrant_store or QdrantKnowledgeStore()
    collection_name = qdrant_collection_for_slug(source.collection_slug)

    # Delete chunks com matching rag_source_id no payload.
    from qdrant_client.models import FieldCondition, Filter, MatchValue

    try:
        await qdrant._client.delete(
            collection_name=collection_name,
            points_selector=Filter(must=[FieldCondition(
                key="rag_source_id",
                match=MatchValue(value=str(source_id)),
            )]),
        )
    except Exception as exc:
        # Colecao pode nao existir mais — ainda assim remove o DB record.
        logger.warning("delete chunks Qdrant falhou: %s", exc)

    await session.delete(source)
    await session.flush()
    return True
